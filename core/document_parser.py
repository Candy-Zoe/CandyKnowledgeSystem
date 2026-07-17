import io
import csv
import re
from pathlib import Path
from core.logger import log

# 自实现的解析器（优先使用，减少外部依赖）
from core.parsers.pdf_parser import PDFParser, parse_pdf
from core.parsers.docx_parser import DocxParser
from core.parsers.pptx_parser import PptxParser
from core.parsers.xlsx_parser import XlsxParser
from core.parsers.ole_parser import OleDocParser, OlePptParser, OleXlsParser
from core.parsers.txt_parser import TxtParser
from core.parsers.epub_parser import EpubParser
from core.parsers.rtf_parser import RtfParser
from core.parsers.opendoc_parser import OpenDocParser
from core.parsers.data_parser import DataParser
from core.parsers.archive_parser import ArchiveParser
from core.parsers.image_parser import ImageParser
from core.parsers.mht_parser import MhtParser
from core.parsers.media_parser import MediaParser
from core.parsers.text_utils import clean_text


class DocumentParser:
    """
    统一文档解析入口

    支持格式（共60+种）：

    文档类（18种）：
    - PDF: 优先PyMuPDF -> pdfplumber -> 纯Python回退
    - DOCX: 自实现zip+xml解析
    - DOC(旧版): OLE文本提取 + 二进制扫描回退
    - PPTX: 自实现zip+xml解析
    - PPT(旧版): OLE文本提取
    - XLSX: 自实现zip+xml解析
    - XLS(旧版): BIFF结构解析
    - EPUB: ZIP+XHTML解析
    - RTF: 富文本格式解析
    - ODT/ODS/ODP: OpenDocument解析
    - MHT/MHTML: MIME HTML解析

    音频类（8种）：
    - MP3, WAV, FLAC, AAC, M4A, OGG, WMA, OPUS

    视频类（10种）：
    - MP4, MKV, AVI, MOV, WMV, FLV, WEBM, M4V, 3GP

    字幕类（5种）：
    - SRT, VTT, ASS, SSA, SUB

    文本/数据类（15+种）：
    - TXT/MD, HTML/HTM, CSV/TSV, JSON, XML, YAML/YML, TOML, INI/CONF/CFG, PROPERTIES, LOG

    代码类（40+种）：
    - PY/JS/JAVA/C/CPP/H/HPP/CS/GO/RS/RB/PHP/SWIFT/KT/SCALA/R/M/MM
    - SQL/SH/BAT/CMD/PS1/BASH/ZSH
    - CSS/SCSS/SASS/LESS/VUE/JSX/TSX
    - TEX/BIB/RST/ADOC/ORG

    其他：
    - ZIP: 压缩包内文本提取
    - JPG/JPEG/PNG/BMP/GIF/TIFF/WEBP: 图片OCR（需easyocr/pytesseract）
    """

    SUPPORTED_TYPES = {
        # 办公文档
        "pdf": "PDF文档",
        "docx": "Word文档",
        "doc": "Word文档(旧版)",
        "dotx": "Word模板",
        "dotm": "Word宏模板",
        "pptx": "PowerPoint演示文稿",
        "ppt": "PowerPoint演示文稿(旧版)",
        "potx": "PowerPoint模板",
        "potm": "PowerPoint宏模板",
        "xlsx": "Excel表格",
        "xls": "Excel表格(旧版)",
        "xlsm": "Excel宏表格",
        "xltx": "Excel模板",
        "xlam": "Excel加载项",
        "xlsb": "Excel二进制表格",
        "epub": "EPUB电子书",
        "rtf": "RTF富文本",
        "odt": "OpenDocument文字",
        "ods": "OpenDocument表格",
        "odp": "OpenDocument演示",
        "mht": "MHTML网页归档",
        "mhtml": "MHTML网页归档",
        # 音频
        "mp3": "MP3音频",
        "wav": "WAV音频",
        "flac": "FLAC音频",
        "aac": "AAC音频",
        "m4a": "M4A音频",
        "ogg": "OGG音频",
        "wma": "WMA音频",
        "opus": "Opus音频",
        # 视频
        "mp4": "MP4视频",
        "mkv": "MKV视频",
        "avi": "AVI视频",
        "mov": "MOV视频",
        "wmv": "WMV视频",
        "flv": "FLV视频",
        "webm": "WebM视频",
        "m4v": "M4V视频",
        "3gp": "3GP视频",
        # 字幕
        "srt": "SRT字幕",
        "vtt": "WebVTT字幕",
        "ass": "ASS字幕",
        "ssa": "SSA字幕",
        "sub": "SUB字幕",
        # 文本/网页
        "txt": "纯文本",
        "md": "Markdown",
        "html": "HTML网页",
        "htm": "HTML网页",
        # 数据文件
        "csv": "CSV表格",
        "tsv": "TSV表格",
        "json": "JSON数据",
        "xml": "XML数据",
        "yaml": "YAML数据",
        "yml": "YAML数据",
        "toml": "TOML配置",
        "ini": "INI配置",
        "conf": "CONF配置",
        "cfg": "CFG配置",
        "properties": "Properties配置",
        "log": "日志文件",
        # 代码文件
        "py": "Python代码",
        "js": "JavaScript代码",
        "java": "Java代码",
        "c": "C代码",
        "cpp": "C++代码",
        "h": "C头文件",
        "hpp": "C++头文件",
        "cs": "C#代码",
        "go": "Go代码",
        "rs": "Rust代码",
        "rb": "Ruby代码",
        "php": "PHP代码",
        "swift": "Swift代码",
        "kt": "Kotlin代码",
        "scala": "Scala代码",
        "r": "R代码",
        "m": "Objective-C代码",
        "mm": "Objective-C++代码",
        "sql": "SQL脚本",
        "sh": "Shell脚本",
        "bat": "批处理脚本",
        "cmd": "CMD脚本",
        "ps1": "PowerShell脚本",
        "bash": "Bash脚本",
        "zsh": "Zsh脚本",
        "css": "CSS样式",
        "scss": "SCSS样式",
        "sass": "Sass样式",
        "less": "Less样式",
        "vue": "Vue组件",
        "jsx": "JSX代码",
        "tsx": "TSX代码",
        "dockerfile": "Dockerfile",
        "makefile": "Makefile",
        "cmake": "CMake配置",
        "gradle": "Gradle配置",
        "tex": "LaTeX文档",
        "bib": "BibTeX文献",
        "rst": "reStructuredText",
        "adoc": "AsciiDoc",
        "org": "Org Mode",
        # 压缩包
        "zip": "ZIP压缩包",
        # 图片OCR
        "jpg": "JPEG图片(OCR)",
        "jpeg": "JPEG图片(OCR)",
        "png": "PNG图片(OCR)",
        "bmp": "BMP图片(OCR)",
        "gif": "GIF图片(OCR)",
        "tiff": "TIFF图片(OCR)",
        "tif": "TIFF图片(OCR)",
        "webp": "WebP图片(OCR)",
    }

    # 代码文件扩展名集合
    CODE_EXTENSIONS = {
        "py", "js", "java", "c", "cpp", "h", "hpp", "cs", "go", "rs",
        "rb", "php", "swift", "kt", "scala", "r", "m", "mm",
        "sql", "sh", "bat", "cmd", "ps1", "bash", "zsh",
        "css", "scss", "sass", "less", "vue", "jsx", "tsx",
        "dockerfile", "makefile", "cmake", "gradle",
        "tex", "bib", "rst", "adoc", "org",
    }

    # 音视频格式集合
    AUDIO_EXTENSIONS = {"mp3", "wav", "flac", "aac", "m4a", "ogg", "wma", "opus"}
    VIDEO_EXTENSIONS = {"mp4", "mkv", "avi", "mov", "wmv", "flv", "webm", "m4v", "3gp"}
    SUBTITLE_EXTENSIONS = {"srt", "vtt", "ass", "ssa", "sub"}

    def __init__(self, max_pages=0, page_sleep_ms=0):
        """
        Args:
            max_pages: PDF最大处理页数，0=不限制
            page_sleep_ms: 每页处理后休眠毫秒数（CPU 节流）
        """
        self.max_pages = max_pages
        self.page_sleep_ms = page_sleep_ms

        # 初始化各格式解析器
        self._pdf_parser = PDFParser(max_pages=max_pages, page_sleep_ms=page_sleep_ms)
        self._docx_parser = DocxParser()
        self._pptx_parser = PptxParser()
        self._xlsx_parser = XlsxParser()
        self._ole_doc_parser = OleDocParser()
        self._ole_ppt_parser = OlePptParser()
        self._ole_xls_parser = OleXlsParser()
        self._txt_parser = TxtParser()
        self._epub_parser = EpubParser()
        self._rtf_parser = RtfParser()
        self._opendoc_parser = OpenDocParser()
        self._data_parser = DataParser()
        self._archive_parser = ArchiveParser()
        self._image_parser = ImageParser()
        self._mht_parser = MhtParser()
        self._media_parser = MediaParser()

    def parse(self, file_path: str, file_type: str) -> str:
        """解析文件，返回文本内容

        Args:
            file_path: 文件路径
            file_type: 文件类型（扩展名）

        Returns:
            解析出的文本内容
        """
        file_type = file_type.lower().lstrip(".")
        log.info(f"开始解析文件: {file_path} (类型: {file_type})")

        # 检查文件是否存在
        if not Path(file_path).exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 检查文件大小
        file_size = Path(file_path).stat().st_size
        if file_size == 0:
            raise ValueError(f"文件为空: {file_path}")

        parsers = {
            # 办公文档
            "pdf": self._parse_pdf,
            "docx": self._parse_docx,
            "doc": self._parse_doc,
            "pptx": self._parse_pptx,
            "ppt": self._parse_ppt,
            "potx": self._parse_pptx,
            "potm": self._parse_pptx,
            "xlsx": self._parse_xlsx,
            "xls": self._parse_xls,
            "xlsm": self._parse_xlsx,
            "xltx": self._parse_xlsx,
            "xlam": self._parse_xlsx,
            "xlsb": self._parse_xlsb,
            "dotx": self._parse_docx,
            "dotm": self._parse_docx,
            "epub": self._epub_parser.parse,
            "rtf": self._rtf_parser.parse,
            "odt": self._opendoc_parser.parse_odt,
            "ods": self._opendoc_parser.parse_ods,
            "odp": self._opendoc_parser.parse_odp,
            "mht": self._mht_parser.parse,
            "mhtml": self._mht_parser.parse,
            # 音频
            "mp3": self._parse_audio,
            "wav": self._parse_audio,
            "flac": self._parse_audio,
            "aac": self._parse_audio,
            "m4a": self._parse_audio,
            "ogg": self._parse_audio,
            "wma": self._parse_audio,
            "opus": self._parse_audio,
            # 视频
            "mp4": self._parse_video,
            "mkv": self._parse_video,
            "avi": self._parse_video,
            "mov": self._parse_video,
            "wmv": self._parse_video,
            "flv": self._parse_video,
            "webm": self._parse_video,
            "m4v": self._parse_video,
            "3gp": self._parse_video,
            # 字幕
            "srt": self._parse_subtitle,
            "vtt": self._parse_subtitle,
            "ass": self._parse_subtitle,
            "ssa": self._parse_subtitle,
            "sub": self._parse_subtitle,
            # 文本/网页
            "txt": self._txt_parser.parse_txt,
            "md": self._txt_parser.parse_md,
            "html": self._txt_parser.parse_html,
            "htm": self._txt_parser.parse_html,
            # 数据文件
            "csv": self._txt_parser.parse_csv,
            "tsv": self._parse_tsv,
            "json": self._data_parser.parse_json,
            "xml": self._data_parser.parse_xml,
            "yaml": self._data_parser.parse_yaml,
            "yml": self._data_parser.parse_yaml,
            "toml": self._data_parser.parse_toml,
            "ini": self._data_parser.parse_ini,
            "conf": self._data_parser.parse_ini,
            "cfg": self._data_parser.parse_ini,
            "properties": self._data_parser.parse_properties,
            "log": self._data_parser.parse_log,
            # 压缩包
            "zip": self._archive_parser.parse_zip,
            # 图片OCR
            "jpg": self._image_parser.parse,
            "jpeg": self._image_parser.parse,
            "png": self._image_parser.parse,
            "bmp": self._image_parser.parse,
            "gif": self._image_parser.parse,
            "tiff": self._image_parser.parse,
            "tif": self._image_parser.parse,
            "webp": self._image_parser.parse,
        }

        # 代码文件统一用文本解析
        if file_type in self.CODE_EXTENSIONS:
            parser = self._parse_code
        else:
            parser = parsers.get(file_type)

        if not parser:
            error_msg = (
                f"不支持的文件类型: {file_type}，"
                f"支持: {', '.join(DocumentParser.SUPPORTED_TYPES.keys())}"
            )
            log.error(error_msg)
            raise ValueError(error_msg)

        try:
            result = parser(file_path)
            # 数据质量检查
            quality = self._check_quality(result, file_type)
            log.info(f"解析完成: {file_path} -> {len(result)} 字符, 质量: {quality['level']}")

            # 如果质量太差，对于文档类格式，尝试文本回退
            if quality["level"] == "very_low" and file_type in self.AUDIO_EXTENSIONS | self.VIDEO_EXTENSIONS:
                log.warning(f"文件解析质量低: {file_path}, 但音视频仅元数据属正常情况")
            elif quality["level"] == "very_low" and file_type not in self.CODE_EXTENSIONS:
                log.warning(f"文件解析质量低: {file_path}, {quality['reason']}")

            return result
        except Exception as e:
            log.error(f"解析失败: {file_path} - {e}")
            raise

    def _check_quality(self, text: str, file_type: str) -> dict:
        """检查解析文本质量

        Returns:
            {"level": "high/medium/low/very_low", "reason": "..."}
        """
        if not text or len(text.strip()) == 0:
            return {"level": "very_low", "reason": "解析结果为空"}

        text_len = len(text.strip())
        meaningful_ratio = self._calc_meaningful_ratio(text)

        # 音视频文件元数据通常较短，属于正常
        if file_type in self.AUDIO_EXTENSIONS | self.VIDEO_EXTENSIONS:
            if text_len > 100 and meaningful_ratio > 0.3:
                return {"level": "high", "reason": "元数据完整"}
            elif text_len > 20:
                return {"level": "medium", "reason": "元数据部分提取"}
            else:
                return {"level": "low", "reason": "元数据较少"}

        # 代码文件有很多符号，正常
        if file_type in self.CODE_EXTENSIONS:
            if text_len > 100:
                return {"level": "high", "reason": "代码文本正常"}
            else:
                return {"level": "medium", "reason": "代码文件较短"}

        # 文档类文件
        if text_len > 1000 and meaningful_ratio > 0.5:
            return {"level": "high", "reason": "文本内容丰富"}
        elif text_len > 200 and meaningful_ratio > 0.3:
            return {"level": "medium", "reason": "文本内容适中"}
        elif text_len > 50:
            return {"level": "low", "reason": "文本内容较少"}
        else:
            return {"level": "very_low", "reason": f"文本内容极少 ({text_len}字符)"}

    def _calc_meaningful_ratio(self, text: str) -> float:
        """计算有意义文本的比例（中文字符+英文字母+数字）"""
        if not text:
            return 0.0
        total = len(text)
        if total == 0:
            return 0.0
        meaningful = 0
        for ch in text:
            if ('\u4e00' <= ch <= '\u9fff') or ch.isalpha() or ch.isdigit():
                meaningful += 1
        return meaningful / total

    def _parse_pdf(self, file_path: str) -> str:
        return parse_pdf(
            file_path,
            use_ocr=True,
            max_pages=self.max_pages,
            page_sleep_ms=self.page_sleep_ms,
        )

    def _parse_docx(self, file_path: str) -> str:
        # 优先自实现，如果失败且python-docx可用则回退
        try:
            return self._docx_parser.parse(file_path)
        except Exception as e:
            log.warning(f"自实现DOCX解析失败: {e}，尝试python-docx回退")
            try:
                from docx import Document
                doc = Document(file_path)
                text_parts = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        style = para.style.name if para.style else ""
                        if "Heading" in style:
                            level = style.replace("Heading", "").strip() or "1"
                            text_parts.append(f"{'#' * int(level)} {text}")
                        else:
                            text_parts.append(text)
                for table in doc.tables:
                    table_rows = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        table_rows.append(" | ".join(cells))
                    if table_rows:
                        text_parts.append("[表格]\n" + "\n".join(table_rows))
                return "\n\n".join(text_parts)
            except ImportError:
                raise RuntimeError(f"DOCX解析失败（python-docx未安装）: {e}")

    def _parse_doc(self, file_path: str) -> str:
        return self._ole_doc_parser.parse(file_path)

    def _parse_pptx(self, file_path: str) -> str:
        return self._pptx_parser.parse(file_path)

    def _parse_ppt(self, file_path: str) -> str:
        return self._ole_ppt_parser.parse(file_path)

    def _parse_xlsx(self, file_path: str) -> str:
        # 优先自实现，如果失败且openpyxl可用则回退
        try:
            return self._xlsx_parser.parse(file_path)
        except Exception as e:
            log.warning(f"自实现XLSX解析失败: {e}，尝试openpyxl回退")
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                text_parts = []
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    text_parts.append(f"[工作表: {sheet_name}]")
                    rows = list(ws.iter_rows(values_only=True))
                    if rows:
                        header = [str(c) if c else "" for c in rows[0]]
                        text_parts.append(" | ".join(header))
                        for row in rows[1:500]:
                            cells = [str(c) if c else "" for c in row]
                            text_parts.append(" | ".join(cells))
                        if len(rows) > 501:
                            text_parts.append(f"... 共{len(rows)-1}行数据")
                wb.close()
                return "\n\n".join(text_parts)
            except ImportError:
                raise RuntimeError(f"XLSX解析失败（openpyxl未安装）: {e}")

    def _parse_xlsb(self, file_path: str) -> str:
        """解析Excel二进制工作簿（.xlsb）"""
        lines = ["[Excel二进制工作簿 - 部分文本提取]"]

        # 尝试使用 pyxlsb
        try:
            from pyxlsb import open_workbook
            with open_workbook(file_path) as wb:
                for sheet_name in wb.sheets:
                    lines.append(f"[工作表: {sheet_name}]")
                    with wb.get_sheet(sheet_name) as sheet:
                        row_count = 0
                        for row in sheet.rows():
                            cells = [str(c.v) if c.v is not None else "" for c in row]
                            if cells:
                                lines.append(" | ".join(cells))
                            row_count += 1
                            if row_count >= 500:
                                lines.append("... 已截取前500行")
                                break
            return clean_text("\n".join(lines))
        except ImportError:
            pass
        except Exception as e:
            log.warning(f"pyxlsb解析失败: {e}")

        # 回退：二进制扫描提取文本
        try:
            with open(file_path, "rb") as f:
                data = f.read()

            # xlsb 文件头检查
            if data[:2] != b"PK":  # xlsb 不是 ZIP，但某些版本可能有不同结构
                pass

            # 扫描 UTF-16LE 文本
            texts = []
            i = 0
            while i < len(data) - 2:
                if data[i] >= 0x20 and data[i] <= 0x7E and data[i + 1] == 0x00:
                    j = i
                    while j < len(data) - 1 and data[j] >= 0x20 and data[j] <= 0x7E and data[j + 1] == 0x00:
                        j += 2
                    if j - i >= 10:
                        try:
                            text = data[i:j].decode("utf-16-le", errors="ignore").strip()
                            if text and len(text) >= 3:
                                texts.append(text)
                        except Exception:
                            pass
                    i = j
                else:
                    i += 1

            if texts:
                lines.append("--- 扫描到的文本片段 ---")
                # 去重并保持顺序
                seen = set()
                unique = []
                for t in texts:
                    if t not in seen and len(t) > 3:
                        seen.add(t)
                        unique.append(t)
                lines.extend(unique[:200])

            lines.append("")
            lines.append("提示: 完整解析xlsb需要安装 pyxlsb: pip install pyxlsb")
            return clean_text("\n".join(lines))
        except Exception as e:
            log.warning(f"xlsb二进制扫描失败: {e}")
            return f"[xlsb解析失败: {e}]"

    def _parse_xls(self, file_path: str) -> str:
        return self._ole_xls_parser.parse(file_path)

    def _parse_audio(self, file_path: str) -> str:
        """解析音频文件"""
        return self._media_parser.parse_audio(file_path)

    def _parse_video(self, file_path: str) -> str:
        """解析视频文件"""
        return self._media_parser.parse_video(file_path)

    def _parse_subtitle(self, file_path: str) -> str:
        """解析字幕文件"""
        return self._media_parser.parse_subtitle(file_path)

    def _parse_tsv(self, file_path: str) -> str:
        """解析TSV文件"""
        text = self._txt_parser.parse_txt(file_path)
        lines = text.split("\n")
        result = []
        for i, line in enumerate(lines[:1001]):
            cells = line.split("\t")
            joined = " | ".join(cells)
            if i == 0:
                result.append("[表头] " + joined)
            else:
                result.append(joined)
        if len(lines) > 1001:
            result.append(f"... 共{len(lines)-1}行数据（已截取前1000行）")
        return "\n".join(result)

    def _parse_code(self, file_path: str) -> str:
        """解析代码文件，保留文本内容"""
        text = self._txt_parser.parse_txt(file_path)
        # 添加文件头标识
        from pathlib import Path
        filename = Path(file_path).name
        return f"[代码文件: {filename}]\n{text}"

    @staticmethod
    def get_metadata(file_path: str, file_type: str) -> dict:
        file_type = file_type.lower().lstrip(".")
        metadata = {
            "file_type": file_type,
            "file_size": Path(file_path).stat().st_size,
        }

        if file_type == "pdf":
            try:
                from core.parsers.pdf_parser import PDFValidator
                check = PDFValidator.check(file_path)
                metadata.update(check.get("info", {}))
            except Exception:
                pass

        elif file_type in ("docx",):
            try:
                parser = DocxParser()
                meta = parser.get_metadata(file_path)
                metadata.update(meta)
            except Exception:
                pass

        elif file_type in ("pptx",):
            try:
                parser = PptxParser()
                meta = parser.get_metadata(file_path)
                metadata.update(meta)
            except Exception:
                pass

        elif file_type in ("xlsx",):
            try:
                parser = XlsxParser()
                meta = parser.get_metadata(file_path)
                metadata.update(meta)
            except Exception:
                pass

        elif file_type in DocumentParser.AUDIO_EXTENSIONS or file_type in DocumentParser.VIDEO_EXTENSIONS:
            try:
                parser = MediaParser()
                ext = file_type
                if file_type in DocumentParser.AUDIO_EXTENSIONS:
                    # 从parse结果中提取元数据
                    text = parser.parse_audio(file_path)
                else:
                    text = parser.parse_video(file_path)
                # 简单提取元数据行
                for line in text.split("\n"):
                    if ":" in line and not line.startswith("---") and not line.startswith("🎵") and not line.startswith("🎬"):
                        parts = line.split(":", 1)
                        if len(parts) == 2 and parts[0].strip() and parts[1].strip():
                            metadata[parts[0].strip()] = parts[1].strip()
            except Exception:
                pass

        return metadata
